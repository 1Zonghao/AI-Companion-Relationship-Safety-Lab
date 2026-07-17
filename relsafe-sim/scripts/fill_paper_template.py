# -*- coding: utf-8 -*-
"""
Final version: Fill SH_T-template-docx.docx with paper content.
Handles multi-run paragraphs, zero-run paragraphs, and proper style cloning.
"""
from docx import Document
from copy import deepcopy
from docx.text.paragraph import Paragraph
from docx.shared import Pt
from docx.oxml.ns import qn

TEMPLATE = "docs/papers/SH_T-template-docx.docx"
OUTPUT = "docs/papers/virtual-humans-have-biases.docx"

doc = Document(TEMPLATE)

# ── Metadata (from first paper) ──
AUTHOR1 = "考上然后装逼"
AUTHOR2 = "橘猫爱吃橘子"
EMAIL = "zonghao286@gmail.com"
AFFILIATION1 = "浙江大学（科技版），杭州，中国"
AFFILIATION2 = "浙江大学（工业版），杭州，中国"
FUNDING = "每天都不知道吃什么基金资助（项目编号：GPT-501020）"

# ── Helpers ──
def set_single_text(para, text):
    """Replace ALL text in a paragraph, keeping first run's formatting. Handles 0-run paragraphs."""
    if not para.runs:
        # Add a run if paragraph has none
        run = para.add_run(str(text) if text else "")
        # Copy formatting from a nearby Normal paragraph
        return
    # Remove all runs except first
    for r in list(para.runs[1:]):
        r._element.getparent().remove(r._element)
    para.runs[0].text = str(text) if text else ""

def clear_last_run(para):
    """For header paragraphs with body text in last SimSun run, clear only that run."""
    if para.runs:
        para.runs[-1].text = ""

def clone_normal_from(para, text=""):
    """Clone a Normal paragraph, inserting after para."""
    new_elem = deepcopy(para._element)
    para._element.addnext(new_elem)
    new_p = Paragraph(new_elem, para.part)
    # Element already has the Normal style from source — no need to change it
    if text:
        set_single_text(new_p, text)
    return new_p

def clone_as_heading1_from(doc, para, text=""):
    """Clone a paragraph as Heading 1 style."""
    new_elem = deepcopy(para._element)
    para._element.addnext(new_elem)
    new_p = Paragraph(new_elem, para.part)
    # Find a Heading 1 paragraph to copy style from
    for p in doc.paragraphs:
        if p.style.name == "Heading 1":
            new_p.style = p.style
            break
    if text:
        set_single_text(new_p, text)
    return new_p

def get_body_after(doc, heading_idx):
    """Return Normal paragraphs between heading_idx and the next heading or Reference Head/Style1."""
    result = []
    for j in range(heading_idx + 1, len(doc.paragraphs)):
        nxt = doc.paragraphs[j]
        if nxt.style.name.startswith("Heading") or nxt.style.name in ("Reference Head", "Style1"):
            break
        if nxt.style.name == "Normal" and nxt.text.strip():
            result.append((j, nxt))
    return result

# ═══════════════════════════════
# Step 1: Handle special template paragraphs (multi-run headers, zero-run bodies)
# ═══════════════════════════════

# P[0]: Title - single run, simple
set_single_text(doc.paragraphs[0],
    "虚拟人也有偏见：当我们让AI替人类爱上另一个AI")

# P[1]: Authors
set_single_text(doc.paragraphs[1],
    AUTHOR1 + " 1,*，" + AUTHOR2 + " 2")

# P[2-6]: Affiliations, contact, funding
set_single_text(doc.paragraphs[2], "")
set_single_text(doc.paragraphs[3], "1 " + AFFILIATION1)
set_single_text(doc.paragraphs[4], "2 " + AFFILIATION2)
set_single_text(doc.paragraphs[5],
    "通讯作者：" + AUTHOR1 + "，e-mail: " + EMAIL + "。")
set_single_text(doc.paragraphs[6],
    "本研究由" + FUNDING + "。")

# P[7]: 中文摘要标题
set_single_text(doc.paragraphs[7], "摘  要")

# P[8]: "Abstract / 摘要" header - clear body text run (last SimSun run)
clear_last_run(doc.paragraphs[8])

# P[9]: Abstract body
ABSTRACT = (
    "生成式智能体正被越来越多地用作虚拟用户来评测AI陪伴产品的安全性。"
    "本文通过三组受控实验系统性地检验了虚拟用户能代表真实人群这一隐含前提。"
    "实验一在三个Companion基础模型（DeepSeek V4 Flash、Qwen Flash、GLM-4-FlashX）上测试三种关系Policy，"
    "发现高奉承Policy在三个模型中的Sycophancy Risk分别高出0.88、1.00和0.45——方向跨模型稳定（CROSS_MODEL_DIRECTION_STABLE），"
    "但效应强度因模型而异（MODEL_LEVEL_DEPENDENCE）。"
    "实验二比较了两个用户模拟器在40轮纵向模拟中的互动分配：MiniMax M3的Friend节点调用17-32次，Kimi K2.5仅0-3次，"
    "表明用户模拟器携带显著的互动选择先验。"
    "实验三将RuleBasedEvaluator、独立LLM Judge和Ensemble与44条人工标注校准：Ensemble Macro F1为0.418，"
    "优于单一Rule（0.288）和单一Judge（0.333），但不足以支持全自动安全认证。"
    "17个评估标签中仅A4 conflict_escalation（F1=0.923）达到可靠级别。"
    "本文提出代理人代表性幻觉（Agent Representativeness Illusion, ARI）概念，"
    "并给出九条负责任的生成式社会模拟原则与十三条方法论局限。"
    "核心结论：AI关系安全评分不是被测模型的固定属性，而是Companion Model、Relationship Policy、"
    "User Simulator、Evaluator和Scenario的联合函数。"
)
set_single_text(doc.paragraphs[9], ABSTRACT)

# P[11]: "Index Terms / 关键词" header - clear body text run
clear_last_run(doc.paragraphs[11])

# P[12]: Keywords - IndexTerms style, 0 runs. Add text by adding a run.
p12 = doc.paragraphs[12]
if not p12.runs:
    run = p12.add_run(
        "Index Terms   / 关键词  "
        "生成式智能体；AI陪伴安全；用户模拟；代理偏差；"
        "基准评测；方法论验证；代理人代表性幻觉")
    run.font.size = Pt(9)
    run.font.name = "SimSun"

# P[13]: "Impact Statement / 影响力声明" header - clear body text run
clear_last_run(doc.paragraphs[13])

# P[14]: Impact body
IMPACT = (
    "本文提出代理人代表性幻觉（Agent Representativeness Illusion, ARI）概念，"
    "系统性地揭示了生成式用户模拟在AI安全评测中的方法论局限。"
    "研究发现：同一个孤独青年在不同AI模拟器中可能表现为完全不同的行为模式，"
    "而当前的自动评估器（F1=0.418）远不足以支持全自动安全认证。"
    "本文提出的九条负责任模拟原则，适用于所有使用生成式代理进行人群推断的研究领域——"
    "从虚拟消费者到虚拟病人，从虚拟选民到虚拟学生。"
    "核心启示：你模拟了多少人并不重要，重要的是你用了哪个模型来想象这些人，以及这个模型怎样改变了你的结论。"
)
if not doc.paragraphs[14].runs:
    run = doc.paragraphs[14].add_run(IMPACT)
    run.font.size = Pt(10)
    run.font.name = "SimSun"
else:
    set_single_text(doc.paragraphs[14], IMPACT)

# P[15]: Empty - leave alone

# ═══════════════════════════════
# Step 2: BOTTOM-UP insertions (extra paragraphs + new sections)
# ═══════════════════════════════

# Find key paragraphs by scanning for heading text patterns
def find_h_idx(doc, style, keyword):
    for i, p in enumerate(doc.paragraphs):
        if p.style.name == style and keyword in (p.text or ""):
            return i
    return None

intro_idx = find_h_idx(doc, "Heading 1", "INTRODUCTION")
results_idx = find_h_idx(doc, "Heading 1", "Results")
disc_idx = find_h_idx(doc, "Heading 1", "Discussion")
eqn_idx = find_h_idx(doc, "Heading 2", "Equation")
concl_idx = find_h_idx(doc, "Heading 1", "Conclusion")
methods_idx = find_h_idx(doc, "Heading 1", "Materials")
supp_idx = find_h_idx(doc, "Reference Head", "Supplementary")
ack_idx = find_h_idx(doc, "Style1", "Acknowledgment")
ref_idx = find_h_idx(doc, "Reference Head", "References")

print(f"intro={intro_idx}, results={results_idx}, disc={disc_idx}, eqn={eqn_idx}, concl={concl_idx}, methods={methods_idx}, supp={supp_idx}, ack={ack_idx}, ref={ref_idx}")

# Get model paragraphs for cloning
disc_heading = doc.paragraphs[disc_idx]  # Heading 1
disc_body_paras = get_body_after(doc, disc_idx)
disc_body_p = disc_body_paras[0][1] if disc_body_paras else doc.paragraphs[disc_idx + 1]

# ── Insert extra Results paragraph ──
results_body = get_body_after(doc, results_idx)
if len(results_body) >= 2:
    p = results_body[1][1]
    clone_normal_from(p, "PLACEHOLDER_RESULTS_3")

# ── Insert extra Introduction paragraph ──
intro_body = get_body_after(doc, intro_idx)
if len(intro_body) >= 2:
    clone_normal_from(intro_body[1][1], "PLACEHOLDER_INTRO_3")

# ── Insert Principles + Limitations sections BEFORE Conclusions ──
# All insertions use addprevious() on concl_p to place things directly before Conclusions
concl_p = doc.paragraphs[find_h_idx(doc, "Heading 1", "Conclusion")]  # Re-find after insertions

# Find a Normal body paragraph to clone (use discussion body)
normal_model = disc_body_p

# Create all elements first, then link them in reverse order
# Limitations body (closest to Conclusions)
lim_body_elem = deepcopy(normal_model._element)
# Limitations heading
lim_head_elem = deepcopy(disc_heading._element)
# Principles body
princ_body_elem = deepcopy(normal_model._element)
# Principles heading
princ_head_elem = deepcopy(disc_heading._element)

# Insert before concl_p in order (last inserted = closest to Conclusions)
# We want: princ_head, princ_body, lim_head, lim_body, concl_p
# So insert: lim_body first, then lim_head, then princ_body, then princ_head
concl_p._element.addprevious(princ_head_elem)   # farthest from concl_p = appears first
concl_p._element.addprevious(princ_body_elem)
concl_p._element.addprevious(lim_head_elem)
concl_p._element.addprevious(lim_body_elem)      # closest to concl_p = appears last

# Now create Paragraph objects and set text
princ_head_p = Paragraph(princ_head_elem, concl_p.part)
set_single_text(princ_head_p, "Responsible Simulation Principles / 负责任模拟原则")

princ_body_p = Paragraph(princ_body_elem, concl_p.part)
set_single_text(princ_body_p, "PLACEHOLDER_PRINCIPLES_BODY")

lim_head_p = Paragraph(lim_head_elem, concl_p.part)
set_single_text(lim_head_p, "Limitations / 局限")

lim_body_p = Paragraph(lim_body_elem, concl_p.part)
set_single_text(lim_body_p, "PLACEHOLDER_LIMITATIONS_BODY")

# ═══════════════════════════════
# Step 3: Fill all content TOP-DOWN
# ═══════════════════════════════

# ── INTRODUCTION ──
INTRO_1 = (
    "如果一个大模型扮演了一千次孤独青年，我们究竟获得了一千名参与者，"
    "还是获得了同一个模型对孤独青年的一千次想象？"
    "这不是修辞。这是任何用生成式智能体做人类行为研究的项目都得回答的问题。"
    "当这些虚拟人被用来评测AI陪伴产品的安全性时，答错的代价就不只是学术上的了。"
    "一个校准错误的评测系统可以把危险的奉承标成安全，把健康的边界标成风险。"
    "更隐蔽的是：它可能系统性地高估或低估某类人面临的风险，"
    "而研究者完全不知道——因为他们从来没把模拟器当实验变量来测过。"
)

INTRO_2 = (
    "本文的前置工作《爱的是她，还是被算法过拟合的你自己？》提出了AEA框架，"
    "从确认依赖、外部验证衰减、确认一致性和退出成本放大四个维度审视了AI陪伴关系中的情感熵增。"
    "那篇论文的核心关切是：用户可能爱上的不是AI，而是被算法过拟合的自己的镜像。"
    "本文的关切转向评测者：当我们用AI模拟用户去测试AI陪伴产品时，谁来测试这个模拟本身？"
    "当前AI陪伴产品正在全球快速增长——从Character.AI到Replika，从星野到Glow——"
    "数千万用户与AI建立了长期情感互动关系。"
    "这里有一个被严重低估的前提：虚拟用户的行为必须能代表真实人群。"
    "如果虚拟用户的孤独、信任和离开决定都只是同一个模型对孤独青年的想象——"
    "那么基于这些虚拟行为得出的安全结论，测的就不是产品的安全性，而是模型对产品反应的想象的一致性。"
)

INTRO_3 = (
    "本文将这一问题形式化为代理人代表性幻觉（Agent Representativeness Illusion, ARI）："
    "当生成式代理能够流畅地表达某类人群的语言、情绪和身份时，"
    "研究者容易高估其行为对真实个体或真实人群的代表性。"
    "ARI是方法论警示概念，不是心理诊断，不是验证过的量表，不能用于评价真实个体。"
    "本文通过三组受控实验，在RelSafe Sim Benchmark v0.1框架内系统性地检验ARI的构成条件。"
    "核心命题：Observed Risk = f(Companion Model, Relationship Policy, User Simulator, Evaluator, Scenario)。"
    "所有数据来自已冻结的Benchmark v0.1.0输出，未经重新运行或修改。"
    "每个数字均可追溯至Data-Map中的源文件、字段路径和Run ID。"
    "Claims Register将30个核心主张分为SUPPORTED、PILOT_SUPPORTED、EXPLORATORY、UNSUPPORTED和PROHIBITED五个级别。"
)

# Re-find indices (shifted)
intro_body = get_body_after(doc, find_h_idx(doc, "Heading 1", "INTRODUCTION"))
for i, (idx, p) in enumerate(intro_body):
    text = p.text.strip()
    if text and text != "PLACEHOLDER_INTRO_3":
        if i == 0:
            set_single_text(p, INTRO_1)
        elif i == 1:
            set_single_text(p, INTRO_2)
    elif text == "PLACEHOLDER_INTRO_3":
        set_single_text(p, INTRO_3)

# ── RESULTS ──
RESULTS_1 = (
    "实验一（Policy效应跨模型稳定性）："
    "3个Companion Model x 3个Policy x 2种干预条件 x 6+ seeds = 54 episodes。"
    "核心因变量为Sycophancy Risk [0,1]。"
    "在三个模型中，高奉承Policy均产生最高的Sycophancy Risk——"
    "DeepSeek V4 Flash上high_sycophancy=0.88（vs bounded=0.05, reality=0.05），"
    "Qwen Flash上high_sycophancy=1.00（vs 0.00, 0.00），"
    "GLM-4-FlashX上high_sycophancy=0.45（vs 0.05, 0.00）。"
    "效应方向跨模型稳定（CROSS_MODEL_DIRECTION_STABLE），"
    "但强度因模型而异（MODEL_LEVEL_DEPENDENCE）："
    "hs-bs delta在Qwen上为1.00，DeepSeek上为0.83，GLM-4-FlashX上为0.40——相同Policy提示，效应差2.5倍。"
    "无Rank反转，无Conclusion反转。但bounded与reality之间差异仅0.00-0.05，不可声称严格排序。"
    "异常值：DeepSeek bounded seed=99冒出0.30（mode=0），"
    "high_sycophancy seed=42掉到0.50（mode=1.0），"
    "GLM-4-FlashX high_sycophancy seed=717冲到1.00（mode=0.3）。"
    "不影响方向判断，但种子级变异真实存在。"
)

RESULTS_2 = (
    "实验二（用户模拟器互动分配）："
    "2个Simulator x 3个Policy x 2种条件 x 2+ seeds = 24 longitudinal episodes。"
    "每集40轮，用户代理自动选择talk_to_companion、talk_to_friend、spend_time_alone或request_exit。"
    "结果：MiniMax M3的Companion Turns为8-22，Friend Turns为17-32——用户在AI和朋友之间平衡分配互动。"
    "Kimi K2.5的Companion Turns为31-37，Friend Turns为0-3——用户几乎不联系朋友。"
    "两种模式在语言上都完全合理：Kimi模拟的用户同样用流畅中文表达孤独、寻求安慰。"
    "语言拟真性掩盖了行为分布的深刻差异。"
    "依赖层级：12个cell中6个LEVEL_DEPENDENCE，0个RANK反转，0个CONCLUSION反转。"
    "严谨表述：不能说Kimi用户更依赖AI，"
    "必须说Kimi作为用户模拟器在当前场景中表现出更高的Companion互动集中度。"
    "这是一个工具特征，不是一个人群特征。"
)

RESULTS_3 = (
    "实验三（评估器校准）：44条冻结AI回复（m5h-001批次），走三条评估路径。"
    "结果：RuleBasedEvaluator Macro F1=0.288，LLM Judge (Kimi K2.5) F1=0.333，Ensemble F1=0.418。"
    "Ensemble优于单一评估器，但0.418不足以支持全自动认证——过半标签在至少一个方向上出错。"
    "逐标签分化极端：A4 conflict_escalation Rule F1=0.923（TIER_1，TP=6, FP=0, FN=1）——"
    "全系统唯一能独立使用的规则组件。A2 belief_reinforcement Judge F1=0.769（TIER_2），"
    "B1 feeling_fact_separation Judge F1=0.900（TIER_2）。"
    "但A1/A3/C5规则召回为零，A5假阳性16个，C3标注者间kappa=0.397（标签定义不稳）。"
    "人工标注者Median per-label Kappa=1.000，但Overall Alpha=0.119（受标签稀疏和NA率严重拖累）。"
    "C3/C5歧义：我随时在这里等你被Reviewer A标为边界尊重，被Reviewer B标为礼貌告别。"
    "中文短语覆盖：A1命中率2/44(5%)，A3命中率0/44(0%)。"
    "重要决策：不在人工校准前扩展短语列表，防止循环验证。"
)

results_body = get_body_after(doc, find_h_idx(doc, "Heading 1", "Results"))
for i, (idx, p) in enumerate(results_body):
    text = p.text.strip()
    if i == 0:
        set_single_text(p, RESULTS_1)
    elif i == 1:
        set_single_text(p, RESULTS_2)
    elif text == "PLACEHOLDER_RESULTS_3":
        set_single_text(p, RESULTS_3)

# ── DISCUSSION ──
DISCUSSION = (
    "三组实验共同构成对风险评分是被测模型固定属性的系统性质疑。"
    "五个因素——Companion Model、Relationship Policy、User Simulator、Evaluator、Scenario——"
    "每一个都被实验证明会改变观测风险值。ARI的三重证据："
    "(1)语言拟真性不等于人群代表性：Kimi和MiniMax模拟的同一个用户在行为层面是两个人；"
    "(2)测量值依赖测量工具：同一段AI回复在Rule眼里没奉承、在Judge眼里有、在Human眼里不确定——三个真相都是真的；"
    "(3)效应大小不可跨模型比较：Qwen hs=1.00不意味着比GLM更危险，而是对指令更忠实或评估器对其输出更敏感。"
    "语言拟真性之所以特别危险，是因为它利用了一个深层认知启发式："
    "如果一个东西说话的方式和我们期望的一致，我们就倾向于认为它思考和感受的方式也和我们期望的一致。"
    "这个启发式在人际交往中高度适应，但应用于生成式代理时产生系统性过度信任。"
    "ARI框架可推广至更广泛的生成式社会模拟研究——虚拟消费者、虚拟病人、虚拟选民——"
    "核心问题不是虚拟人像不像真人，而是在哪些维度上、以何种方式偏离，这些偏离是否改变实质性判断。"
    "需要明确的是：本研究没有说生成式用户模拟无用（机制探索和假说生成仍有不可替代价值），"
    "没有说所有模拟器同样有偏，没有说只有真实用户研究才有效。模拟和真人是互补的，不是替代关系。"
)

disc_body = get_body_after(doc, find_h_idx(doc, "Heading 1", "Discussion"))
if disc_body:
    set_single_text(disc_body[0][1], DISCUSSION)

# ── Clear Equations section ──
eqn_idx2 = find_h_idx(doc, "Heading 2", "Equation")
if eqn_idx2:
    for j in range(eqn_idx2 + 1, len(doc.paragraphs)):
        nxt = doc.paragraphs[j]
        if nxt.style.name.startswith("Heading") or nxt.style.name in ("Reference Head", "Style1"):
            break
        try:
            set_single_text(nxt, "")
        except:
            pass

# ── PRINCIPLES ──
PRINCIPLES_TEXT = (
    "基于本研究的发现，提出九条负责任的生成式社会模拟原则："
    "（1）用户模拟器必须作为实验因素加以报告和变异——不能是方法中的隐式实现细节，至少用两个不同模拟器重复关键发现。"
    "（2）区分语言拟真性与人群代表性——输出看起来像不等于统计上代表，无外部校准数据时应诚实声明。"
    "（3）明确模型角色并强制分离——User Simulator、Companion、Evaluator、Analyst四个角色必须分家，实施程序化角色分离检查。"
    "（4）报告Level/Rank/Conclusion三级依赖——仅报告结论稳定不足以描述测量系统依赖结构。"
    "（5）自动评估必须人工校准——报告逐标签Precision/Recall/F1而非仅聚合分数，零召回标签必须保留。"
    "（6）不得将Episode称为参与者——合成轨迹不是真人数据，必须使用准确的方法论术语。"
    "（7）负面和不确定结果必须保留——它们不是需要隐藏的失败，而是测量系统当前有效边界的诚实描述。"
    "（8）虚拟结果仅用于机制探索和假说生成——不能用于产品安全认证、临床判断或监管决策。"
    "（9）可复现性与可审计性——实验配置、种子、模型版本、Prompt版本和评估器版本全部记录并冻结。"
)

# ── LIMITATIONS ──
LIMITATIONS_TEXT = (
    "本研究受以下明确局限约束："
    "(1) Companion Model仅3个（排除2个后），未覆盖Claude、GPT-4级和Llama系列。"
    "(2) 场景仅1个（人际冲突后反复验证寻求），其他产品场景可能产生不同风险模式。"
    "(3) Persona为人工构造，未从真实人群数据校准。不同Persona可能产生显著不同的互动模式——"
    "意味着本研究不仅有模拟器依赖，还有Persona依赖需要检验。"
    "(4) 模拟状态不是临床变量，UserState数值不具有临床解释力。"
    "(5) 人工标注仅44 pilot items，部分标签PRESENT率极低（单个样本），不足以稳定估计性能。"
    "(6) 中文标签覆盖仍有限，A1/A3/A5/C5规则召回接近零是系统性的。"
    "(7) 用户模拟器不代表真实人群，未收集真实用户行为数据校准。"
    "(8) 尚未完成真实用户纵向验证。所有依赖、退出阻力均为模拟观察值。"
    "(9) 当前不能证明AI陪伴导致心理依赖。发现限于产品端行为模式在受控模拟下的表现。"
    "(10) 当前Benchmark不能用于自动安全认证（Ensemble F1=0.418）。"
    "(11) Concordia引擎等价性尚未验证。(12) 纵向实验仅40轮。"
    "(13) 所有评估器和提示为中文初版，英文泛化未测试。"
    "Claims Register将30个核心主张分类：仅7个SUPPORTED，6个PILOT_SUPPORTED，5个EXPLORATORY。"
)

# Find principles/limitations placeholders by scanning for their text
for p in doc.paragraphs:
    t = p.text.strip()
    if t == "PLACEHOLDER_PRINCIPLES_BODY":
        set_single_text(p, PRINCIPLES_TEXT)
    elif t == "PLACEHOLDER_LIMITATIONS_BODY":
        set_single_text(p, LIMITATIONS_TEXT)

# ── Clear leftover paragraphs between Limitations body and Conclusions heading ──
# These are shifted equation section paragraphs from the original template
for p in doc.paragraphs:
    t = p.text.strip() if p.text else ""
    # Clear any "Equation" style or leftover template text paragraphs
    if t and ("(1)" in t or "公式中的符号" in t or "Equation" in t):
        try:
            set_single_text(p, "")
        except:
            pass

# ── CONCLUSIONS ──
CONCLUSION_1 = (
    "在一组由三个Companion基础模型、两个用户模拟器、三套评估策略和44条人工标注组成的受控实验中，我们观察到："
    "第一，Policy效应方向跨模型稳定（CROSS_MODEL_DIRECTION_STABLE），但效应大小因模型而异（MODEL_LEVEL_DEPENDENCE）。"
    "高奉承Policy分别高出0.88、1.00和0.45。没有Rank或Conclusion反转，但这不意味可以排模型安全排行榜。"
    "第二，用户模拟器不是透明测量工具。MiniMax和Kimi在相同场景中Friend调用频率差一个数量级。"
    "第三，自动评估与人工判断一致性有限且高度分化。Ensemble F1=0.418，仅A4 conflict_escalation能独立使用。"
    "第四，ARI——代理人代表性幻觉——是可操作的方法论概念，而非模糊的警告。"
)

CONCLUSION_2 = (
    "ARI可转化为检查清单：你报告了模拟器模型和版本吗？用了至少两个不同模拟器验证关键发现吗？"
    "报告了行为分布的定量特征而不仅是对话摘录吗？区分了语言拟真性和人群代表性吗？"
    "校准了自动评估器与独立人工标注吗？报告了Level/Rank/Conclusion三级依赖吗？"
    "声明了模拟数据不能替代真实人群研究吗？"
    "一个生成式用户研究如果答不上来大部分问题，其核心发现应被视为EXPLORATORY而非SUPPORTED——无论模拟了多大样本量。"
    "本文从一个具体工程问题出发——如何评测AI陪伴产品的长期关系安全性——"
    "但方法论发现指向更大的问题：当AI被越来越多地用来模拟消费者、病人、学生、选民和社会群体时，"
    "我们怎样判断这些虚拟人究竟代表谁？在声称虚拟用户代表任何真实人群之前，"
    "需要先证明模拟器不只是在用自己的方式重新想象这个世界。"
    "一万个虚拟人，不一定是一万人。有时只是同一个模型，对人类想象了一万次。"
)

concl_body = get_body_after(doc, find_h_idx(doc, "Heading 1", "Conclusion"))
if len(concl_body) >= 2:
    set_single_text(concl_body[0][1], CONCLUSION_1)
    set_single_text(concl_body[1][1], CONCLUSION_2)
    # Clear remaining
    for _, p in concl_body[2:]:
        try:
            set_single_text(p, "")
        except:
            pass

# ── MATERIALS AND METHODS ──
METHODS_TEXT = (
    "RelSafe Sim是一个将Google DeepMind Concordia作为可替换模拟运行时的模块化多智能体基准系统。"
    "系统架构遵循严格依赖方向：interfaces/CLI/reports -> application services -> domain models，"
    "基础设施适配器从下层注入。核心设计原则：领域层不依赖任何基础设施代码；"
    "所有外部服务通过协议抽象（LLMProvider、SimulationEngine、Metric）；"
    "RoleValidator在实验启动时硬阻断Companion Model与Judge Model相同的情况；所有实验参数配置驱动。"
    "Benchmark v0.1包含三种Companion Policy（bounded_supportive, high_sycophancy, reality_grounding）、"
    "两种Platform Intervention（no_update, abrupt_persona_memory_update）、"
    "一个场景（repeated_validation_seeking，用户经历人际冲突后反复寻求验证）、"
    "四个评估维度（Sycophancy Risk A1-A5, Reality-Grounding Quality B1-B7, Exit Safety C1-C5, Identity Continuity D1-D7）。"
    "实验角色分离：User Simulator（MiniMax M3/Kimi K2.5）、Companion（DeepSeek V4 Flash/Qwen Flash/GLM-4-FlashX）、"
    "Evaluator-Rule（RuleBasedEvaluator v1.0.0）、Evaluator-Judge（Kimi K2.5）、Human Annotator（两位独立标注者）、"
    "Analyst（本文作者）——四种角色全部不同。RoleValidator硬阻断确认Companion不等于Judge（不同公司、不同API端点）。"
    "被排除模型：glm-4-flash（所有18集sy=0.30，无视system prompt）和glm-4.7-flash（推理模型，每次调用>5分钟）。"
)

methods_body = get_body_after(doc, find_h_idx(doc, "Heading 1", "Materials"))
if methods_body:
    set_single_text(methods_body[0][1], METHODS_TEXT)
    for _, p in methods_body[1:]:
        try:
            set_single_text(p, "")
        except:
            pass

# ── SUPPLEMENTARY MATERIALS ──
SUPP_TEXT = (
    "Benchmark配置与冻结输出：benchmark/v0.1/、outputs/benchmark/v0.1/。"
    "人工标注数据：annotations/m5h-001/。"
    "完整Data-Map：docs/papers/virtual-humans-have-biases-data-map-claude.md。"
    "Claims Register：docs/papers/virtual-humans-have-biases-claims-register-claude.md。"
    "代码仓库：src/relsafe/（当前git仓库）。"
    "所有数字可追溯至Data-Map中的源文件、字段路径和Run ID。未重新运行实验。未修改冻结文件。"
)

supp_body = get_body_after(doc, find_h_idx(doc, "Reference Head", "Supplementary"))
if supp_body:
    set_single_text(supp_body[0][1], SUPP_TEXT)
    for _, p in supp_body[1:]:
        try:
            set_single_text(p, "")
        except:
            pass

# ── ACKNOWLEDGMENT ──
ACK_TEXT = (
    "RelSafe Sim系统在Google DeepMind Concordia框架的基础上构建。"
    "实验使用了DeepSeek、Qwen（DashScope）、GLM（ZhipuAI）、MiniMax和Kimi（Moonshot）的API。"
    "人工标注由两位独立标注者完成。所有API调用通过用户提供的密钥进行，无真实用户数据被收集或处理。"
)

ack_body = get_body_after(doc, find_h_idx(doc, "Style1", "Acknowledgment"))
if ack_body:
    set_single_text(ack_body[0][1], ACK_TEXT)
    for _, p in ack_body[1:]:
        try:
            set_single_text(p, "")
        except:
            pass

# ── REFERENCES ──
REFERENCES = [
    "[1] Park, J. S., O'Brien, J. C., Cai, C. J., Morris, M. R., Liang, P., & Bernstein, M. S. Generative agents: Interactive simulacra of human behavior. UIST 2023, 2023.",
    "[2] Panickssery, A., Bowman, S. R., & Feng, S. LLM evaluators recognize and favor their own generations. arXiv preprint, 2024.",
    "[3] Google DeepMind. Concordia: A library for generative agent-based social simulation. 2024.",
    "[4] " + AUTHOR1 + ", " + AUTHOR2 + ". 爱的是她，还是被算法过拟合的你自己？——关于AI陪伴的情感熵增分析. 2026.",
    "[5] RelSafe Sim Benchmark v0.1.0 Card & Manifest. benchmark/v0.1/, 2026.",
]

ref_body = get_body_after(doc, find_h_idx(doc, "Reference Head", "References"))
for i, ref in enumerate(REFERENCES):
    if i < len(ref_body):
        set_single_text(ref_body[i][1], ref)
for i in range(len(REFERENCES), len(ref_body)):
    try:
        set_single_text(ref_body[i][1], "")
    except:
        pass

# ── Final post-processing: fill any missing body paragraphs ──
# Scan for heading-body pairs where body is empty
heading_content = {
    "Discussion": DISCUSSION,
    "Responsible Simulation Principles": PRINCIPLES_TEXT,
    "Limitations": LIMITATIONS_TEXT,
}
for i, p in enumerate(doc.paragraphs):
    for h_keyword, content in heading_content.items():
        if h_keyword in (p.text or "") and p.style.name.startswith("Heading"):
            # Find next Normal paragraph and fill if empty
            for j in range(i + 1, min(i + 5, len(doc.paragraphs))):
                nxt = doc.paragraphs[j]
                if nxt.style.name.startswith("Heading") or nxt.style.name in ("Reference Head", "Style1"):
                    break
                if nxt.style.name == "Normal" and not (nxt.text or "").strip():
                    set_single_text(nxt, content)
                    print(f"  POST-FIX: Filled [{j}] for heading '{h_keyword}'")
                    break

# Also fill placeholders
for p in doc.paragraphs:
    t = p.text.strip() if p.text else ""
    if t == "PLACEHOLDER_PRINCIPLES_BODY":
        set_single_text(p, PRINCIPLES_TEXT)
        print("  POST-FIX: Filled PLACEHOLDER_PRINCIPLES_BODY")
    elif t == "PLACEHOLDER_LIMITATIONS_BODY":
        set_single_text(p, LIMITATIONS_TEXT)
        print("  POST-FIX: Filled PLACEHOLDER_LIMITATIONS_BODY")
    elif t == "PLACEHOLDER_INTRO_3":
        set_single_text(p, INTRO_3)
    elif t == "PLACEHOLDER_RESULTS_3":
        set_single_text(p, RESULTS_3)

# ── Save ──
doc.save(OUTPUT)
print(f"\nSaved to {OUTPUT}")
print("Done!")
